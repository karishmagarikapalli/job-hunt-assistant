import os
import logging
import json
import re
from datetime import datetime
from jinja2 import Template

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(os.path.join(os.path.dirname(__file__), '..', 'logs', 'document_generator.log')),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger('document_generator')

class DocumentGenerator:
    """
    AI-powered document generator for creating tailored resumes and cover letters
    based on job descriptions and user profiles.
    """
    
    def __init__(self, templates_dir=None):
        """
        Initialize the document generator.
        
        Args:
            templates_dir (str): Directory containing document templates
        """
        self.templates_dir = templates_dir or os.path.join(os.path.dirname(__file__), '..', 'templates')
        
        # Create logs directory if it doesn't exist
        os.makedirs(os.path.join(os.path.dirname(__file__), '..', 'logs'), exist_ok=True)
        
        # Create templates directory if it doesn't exist
        os.makedirs(self.templates_dir, exist_ok=True)
        
        # Load templates
        self.resume_templates = self._load_templates('resume')
        self.cover_letter_templates = self._load_templates('cover_letter')
    
    def _load_templates(self, template_type):
        """
        Load templates of the specified type.
        
        Args:
            template_type (str): Type of template ('resume' or 'cover_letter')
        
        Returns:
            dict: Dictionary of templates
        """
        templates = {}
        template_dir = os.path.join(self.templates_dir, template_type)
        
        # Create template type directory if it doesn't exist
        os.makedirs(template_dir, exist_ok=True)
        
        # Load templates from directory
        try:
            for filename in os.listdir(template_dir):
                if filename.endswith('.html') or filename.endswith('.tex') or filename.endswith('.md'):
                    template_id = os.path.splitext(filename)[0]
                    template_path = os.path.join(template_dir, filename)
                    
                    with open(template_path, 'r') as f:
                        template_content = f.read()
                    
                    templates[template_id] = {
                        'id': template_id,
                        'path': template_path,
                        'content': template_content,
                        'format': os.path.splitext(filename)[1][1:]  # Get extension without dot
                    }
            
            logger.info(f"Loaded {len(templates)} {template_type} templates")
        except Exception as e:
            logger.error(f"Error loading {template_type} templates: {str(e)}")
        
        # If no templates found, create default templates
        if not templates:
            logger.info(f"No {template_type} templates found, creating defaults")
            self._create_default_templates(template_type)
            return self._load_templates(template_type)  # Recursive call to load newly created templates
        
        return templates
    
    def _create_default_templates(self, template_type):
        """
        Create default templates of the specified type.
        
        Args:
            template_type (str): Type of template ('resume' or 'cover_letter')
        """
        template_dir = os.path.join(self.templates_dir, template_type)
        
        if template_type == 'resume':
            # Create default resume templates
            templates = {
                'professional': {
                    'name': 'Professional Resume',
                    'content': """<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{{ user.first_name }} {{ user.last_name }} - Resume</title>
    <style>
        body {
            font-family: 'Calibri', 'Arial', sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
        }
        .header {
            text-align: center;
            margin-bottom: 20px;
        }
        h1 {
            font-size: 24px;
            margin-bottom: 5px;
            color: #2a5885;
        }
        .contact-info {
            font-size: 14px;
            margin-bottom: 20px;
        }
        h2 {
            font-size: 18px;
            border-bottom: 1px solid #2a5885;
            padding-bottom: 5px;
            color: #2a5885;
            margin-top: 20px;
        }
        .section {
            margin-bottom: 20px;
        }
        .job {
            margin-bottom: 15px;
        }
        .job-header {
            display: flex;
            justify-content: space-between;
            font-weight: bold;
        }
        .skills {
            display: flex;
            flex-wrap: wrap;
            gap: 10px;
        }
        .skill {
            background-color: #f0f0f0;
            padding: 5px 10px;
            border-radius: 3px;
            font-size: 14px;
        }
        .education-item {
            margin-bottom: 10px;
        }
        ul {
            margin-top: 5px;
        }
        li {
            margin-bottom: 5px;
        }
    </style>
</head>
<body>
    <div class="header">
        <h1>{{ user.first_name }} {{ user.last_name }}</h1>
        <div class="contact-info">
            {{ user.email }} | {{ user.phone }} | {{ user.location }}
            {% if user.linkedin_url %} | <a href="{{ user.linkedin_url }}">LinkedIn</a>{% endif %}
        </div>
    </div>
    
    {% if user.summary %}
    <div class="section">
        <h2>Professional Summary</h2>
        <p>{{ user.summary }}</p>
    </div>
    {% endif %}
    
    <div class="section">
        <h2>Skills</h2>
        <div class="skills">
            {% for skill in user.skills %}
            <span class="skill">{{ skill }}</span>
            {% endfor %}
        </div>
    </div>
    
    <div class="section">
        <h2>Work Experience</h2>
        {% for job in user.work_history %}
        <div class="job">
            <div class="job-header">
                <span>{{ job.title }} at {{ job.company }}</span>
                <span>{{ job.start_date }} - {{ job.end_date or 'Present' }}</span>
            </div>
            <ul>
                {% for achievement in job.achievements %}
                <li>{{ achievement }}</li>
                {% endfor %}
            </ul>
        </div>
        {% endfor %}
    </div>
    
    <div class="section">
        <h2>Education</h2>
        {% for edu in user.education %}
        <div class="education-item">
            <div class="job-header">
                <span>{{ edu.degree }} in {{ edu.field }}</span>
                <span>{{ edu.graduation_year }}</span>
            </div>
            <div>{{ edu.institution }}, {{ edu.location }}</div>
        </div>
        {% endfor %}
    </div>
    
    {% if user.certifications %}
    <div class="section">
        <h2>Certifications</h2>
        <ul>
            {% for cert in user.certifications %}
            <li>{{ cert.name }} ({{ cert.year }})</li>
            {% endfor %}
        </ul>
    </div>
    {% endif %}
</body>
</html>"""
                },
                'modern': {
                    'name': 'Modern Resume',
                    'content': """<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{{ user.first_name }} {{ user.last_name }} - Resume</title>
    <style>
        body {
            font-family: 'Segoe UI', 'Helvetica', sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f9f9f9;
        }
        .container {
            background-color: white;
            padding: 30px;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
        }
        .header {
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 30px;
        }
        .name-title {
            flex: 2;
        }
        .contact-info {
            flex: 1;
            text-align: right;
            font-size: 14px;
        }
        h1 {
            font-size: 28px;
            margin: 0;
            color: #2d3e50;
        }
        .job-title {
            font-size: 18px;
            color: #3498db;
            margin-top: 5px;
        }
        h2 {
            font-size: 20px;
            color: #2d3e50;
            margin-top: 25px;
            padding-bottom: 10px;
            border-bottom: 2px solid #3498db;
        }
        .section {
            margin-bottom: 25px;
        }
        .job {
            margin-bottom: 20px;
        }
        .job-header {
            display: flex;
            justify-content: space-between;
            font-weight: bold;
            margin-bottom: 5px;
        }
        .company {
            color: #3498db;
        }
        .date {
            color: #7f8c8d;
            font-size: 14px;
        }
        .skills-container {
            display: flex;
            flex-wrap: wrap;
            gap: 10px;
            margin-top: 10px;
        }
        .skill {
            background-color: #e8f4fc;
            color: #3498db;
            padding: 5px 12px;
            border-radius: 15px;
            font-size: 14px;
        }
        .education-item {
            margin-bottom: 15px;
        }
        ul {
            margin-top: 8px;
            padding-left: 20px;
        }
        li {
            margin-bottom: 8px;
        }
        a {
            color: #3498db;
            text-decoration: none;
        }
        a:hover {
            text-decoration: underline;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <div class="name-title">
                <h1>{{ user.first_name }} {{ user.last_name }}</h1>
                <div class="job-title">{{ user.current_title }}</div>
            </div>
            <div class="contact-info">
                {{ user.email }}<br>
                {{ user.phone }}<br>
                {{ user.location }}<br>
                {% if user.linkedin_url %}<a href="{{ user.linkedin_url }}">LinkedIn</a>{% endif %}
            </div>
        </div>
        
        {% if user.summary %}
        <div class="section">
            <h2>About Me</h2>
            <p>{{ user.summary }}</p>
        </div>
        {% endif %}
        
        <div class="section">
            <h2>Experience</h2>
            {% for job in user.work_history %}
            <div class="job">
                <div class="job-header">
                    <span class="company">{{ job.title }} | {{ job.company }}</span>
                    <span class="date">{{ job.start_date }} - {{ job.end_date or 'Present' }}</span>
                </div>
                <ul>
                    {% for achievement in job.achievements %}
                    <li>{{ achievement }}</li>
                    {% endfor %}
                </ul>
            </div>
            {% endfor %}
        </div>
        
        <div class="section">
            <h2>Skills</h2>
            <div class="skills-container">
                {% for skill in user.skills %}
                <span class="skill">{{ skill }}</span>
                {% endfor %}
            </div>
        </div>
        
        <div class="section">
            <h2>Education</h2>
            {% for edu in user.education %}
            <div class="education-item">
                <div class="job-header">
                    <span class="company">{{ edu.degree }} in {{ edu.field }}</span>
                    <span class="date">{{ edu.graduation_year }}</span>
                </div>
                <div>{{ edu.institution }}, {{ edu.location }}</div>
            </div>
            {% endfor %}
        </div>
        
        {% if user.certifications %}
        <div class="section">
            <h2>Certifications</h2>
            <ul>
                {% for cert in user.certifications %}
                <li>{{ cert.name }} ({{ cert.year }})</li>
                {% endfor %}
            </ul>
        </div>
        {% endif %}
    </div>
</body>
</html>"""
                },
                'technical': {
                    'name': 'Technical Resume',
                    'content': """<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{{ user.first_name }} {{ user.last_name }} - Technical Resume</title>
    <style>
        body {
            font-family: 'Roboto', 'Arial', sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
        }
        .header {
            text-align: center;
            margin-bottom: 20px;
        }
        h1 {
            font-size: 24px;
            margin-bottom: 5px;
            color: #333;
        }
        .title {
            font-size: 18px;
            color: #666;
            margin-bottom: 10px;
        }
        .contact-info {
            font-size: 14px;
            margin-bottom: 20px;
        }
        h2 {
            font-size: 18px;
            background-color: #f0f0f0;
            padding: 5px 10px;
            color: #333;
            margin-top: 20px;
        }
        .section {
            margin-bottom: 20px;
        }
        .job {
            margin-bottom: 15px;
        }
        .job-header {
            display: flex;
            justify-content: space-between;
            font-weight: bold;
        }
        .skills-grid {
            display: grid;
            grid-template-columns: repeat(3, 1fr);
            gap: 15px;
            margin-top: 10px;
        }
        .skill-category {
            margin-bottom: 15px;
        }
        .skill-category h3 {
            font-size: 16px;
            margin-bottom: 5px;
            color: #555;
        }
        .skill-list {
            list-style-type: none;
            padding-left: 0;
            margin: 0;
        }
        .skill-list li {
            margin-bottom: 3px;
            font-size: 14px;
        }
        .education-item {
            margin-bottom: 10px;
        }
        .project {
            margin-bottom: 15px;
        }
        .project-header {
            font-weight: bold;
            margin-bottom: 5px;
        }
        ul {
            margin-top: 5px;
        }
        li {
            margin-bottom: 5px;
        }
    </style>
</head>
<body>
    <div class="header">
        <h1>{{ user.first_name }} {{ user.last_name }}</h1>
        <div class="title">{{ user.current_title }}</div>
        <div class="contact-info">
            {{ user.email }} | {{ user.phone }} | {{ user.location }}
            {% if user.linkedin_url %} | <a href="{{ user.linkedin_url }}">LinkedIn</a>{% endif %}
            {% if user.github_url %} | <a href="{{ user.github_url }}">GitHub</a>{% endif %}
        </div>
    </div>
    
    {% if user.summary %}
    <div class="section">
        <h2>Technical Profile</h2>
        <p>{{ user.summary }}</p>
    </div>
    {% endif %}
    
    <div class="section">
        <h2>Technical Skills</h2>
        <div class="skills-grid">
            <div class="skill-category">
                <h3>Languages</h3>
                <ul class="skill-list">
                    {% for skill in user.skills %}
                    {% if skill in ['Python', 'JavaScript', 'Java', 'C++', 'C#', 'TypeScript', 'Go', 'Ruby', 'PHP', 'Swift', 'Kotlin'] %}
                    <li>{{ skill }}</li>
                    {% endif %}
                    {% endfor %}
                </ul>
            </div>
            <div class="skill-category">
                <h3>Frameworks & Libraries</h3>
                <ul class="skill-list">
                    {% for skill in user.skills %}
                    {% if skill in ['React', 'Angular', 'Vue.js', 'Django', 'Flask', 'Spring', 'Express', 'Node.js', 'TensorFlow', 'PyTorch'] %}
                    <li>{{ skill }}</li>
                    {% endif %}
                    {% endfor %}
                </ul>
            </div>
            <div class="skill-category">
                <h3>Databases & Cloud</h3>
                <ul class="skill-list">
                    {% for skill in user.skills %}
                    {% if skill in ['SQL', 'MySQL', 'PostgreSQL', 'MongoDB', 'AWS', 'Azure', 'GCP', 'Docker', 'Kubernetes'] %}
                    <li>{{ skill }}</li>
                    {% endif %}
                    {% endfor %}
                </ul>
            </div>
        </div>
    </div>
    
    <div class="section">
        <h2>Professional Experience</h2>
        {% for job in user.work_history %}
        <div class="job">
            <div class="job-header">
                <span>{{ job.title }} at {{ job.company }}</span>
                <span>{{ job.start_date }} - {{ job.end_date or 'Present' }}</span>
            </div>
            <ul>
                {% for achievement in job.achievements %}
                <li>{{ achievement }}</li>
                {% endfor %}
            </ul>
        </div>
        {% endfor %}
    </div>
    
    {% if user.projects %}
    <div class="section">
        <h2>Technical Projects</h2>
        {% for project in user.projects %}
        <div class="project">
            <div class="project-header">{{ project.name }}</div>
            <p>{{ project.description }}</p>
            <p><strong>Technologies:</strong> {{ project.technologies }}</p>
            {% if project.url %}<p><a href="{{ project.url }}">Project Link</a></p>{% endif %}
        </div>
        {% endfor %}
    </div>
    {% endif %}
    
    <div class="section">
        <h2>Education</h2>
        {% for edu in user.education %}
        <div class="education-item">
            <div class="job-header">
                <span>{{ edu.degree }} in {{ edu.field }}</span>
                <span>{{ edu.graduation_year }}</span>
            </div>
            <div>{{ edu.institution }}, {{ edu.location }}</div>
        </div>
        {% endfor %}
    </div>
    
    {% if user.certifications %}
    <div class="section">
        <h2>Certifications</h2>
        <ul>
            {% for cert in user.certifications %}
            <li>{{ cert.name }} ({{ cert.year }})</li>
            {% endfor %}
        </ul>
    </div>
    {% endif %}
</body>
</html>"""
                }
            }
        elif template_type == 'cover_letter':
            # Create default cover letter templates
            templates = {
                'professional': {
                    'name': 'Professional Cover Letter',
                    'content': """<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{{ user.first_name }} {{ user.last_name }} - Cover Letter</title>
    <style>
        body {
            font-family: 'Calibri', 'Arial', sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
        }
        .header {
            margin-bottom: 30px;
        }
        .date {
            margin-bottom: 20px;
        }
        .recipient {
            margin-bottom: 20px;
        }
        .greeting {
            margin-bottom: 20px;
        }
        .content {
            margin-bottom: 30px;
        }
        .closing {
            margin-bottom: 40px;
        }
        .signature {
            font-weight: bold;
        }
    </style>
</head>
<body>
    <div class="header">
        <div>{{ user.first_name }} {{ user.last_name }}</div>
        <div>{{ user.email }}</div>
        <div>{{ user.phone }}</div>
        <div>{{ user.location }}</div>
    </div>
    
    <div class="date">{{ current_date }}</div>
    
    <div class="recipient">
        <div>{{ job.hiring_manager or 'Hiring Manager' }}</div>
        <div>{{ job.company }}</div>
        <div>{{ job.location }}</div>
    </div>
    
    <div class="greeting">Dear {{ job.hiring_manager or 'Hiring Manager' }},</div>
    
    <div class="content">
        <p>I am writing to express my interest in the {{ job.title }} position at {{ job.company }}. With my background in {{ user.background }} and experience in {{ user.experience }}, I am confident that I would be a valuable addition to your team.</p>
        
        <p>Throughout my career, I have developed strong skills in {{ user.key_skills }}. In my current role as {{ user.current_title }} at {{ user.current_company }}, I have {{ user.current_achievements }}. Prior to this, as {{ user.previous_title }} at {{ user.previous_company }}, I successfully {{ user.previous_achievements }}.</p>
        
        <p>I am particularly drawn to {{ job.company }} because {{ job.company_appeal }}. The {{ job.title }} position appeals to me because {{ job.position_appeal }}. I am excited about the opportunity to {{ job.opportunity }}.</p>
        
        <p>I am confident that my skills in {{ job.required_skills }} make me an ideal candidate for this position. My experience with {{ job.relevant_experience }} has prepared me well for the challenges of this role.</p>
        
        <p>Thank you for considering my application. I look forward to the opportunity to discuss how my background, skills, and experiences would benefit {{ job.company }}.</p>
    </div>
    
    <div class="closing">Sincerely,</div>
    
    <div class="signature">{{ user.first_name }} {{ user.last_name }}</div>
</body>
</html>"""
                },
                'modern': {
                    'name': 'Modern Cover Letter',
                    'content': """<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{{ user.first_name }} {{ user.last_name }} - Cover Letter</title>
    <style>
        body {
            font-family: 'Segoe UI', 'Helvetica', sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f9f9f9;
        }
        .container {
            background-color: white;
            padding: 40px;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
        }
        .header {
            text-align: center;
            margin-bottom: 30px;
            padding-bottom: 20px;
            border-bottom: 1px solid #eee;
        }
        .name {
            font-size: 24px;
            font-weight: bold;
            color: #2d3e50;
            margin-bottom: 5px;
        }
        .contact {
            color: #7f8c8d;
            font-size: 14px;
        }
        .date {
            text-align: right;
            margin-bottom: 30px;
            color: #7f8c8d;
        }
        .recipient {
            margin-bottom: 30px;
        }
        .greeting {
            font-size: 18px;
            margin-bottom: 20px;
            color: #2d3e50;
        }
        .content {
            margin-bottom: 30px;
        }
        .content p {
            margin-bottom: 15px;
        }
        .closing {
            margin-bottom: 20px;
        }
        .signature {
            font-weight: bold;
            color: #2d3e50;
        }
        a {
            color: #3498db;
            text-decoration: none;
        }
        a:hover {
            text-decoration: underline;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <div class="name">{{ user.first_name }} {{ user.last_name }}</div>
            <div class="contact">
                {{ user.email }} | {{ user.phone }} | {{ user.location }}
                {% if user.linkedin_url %} | <a href="{{ user.linkedin_url }}">LinkedIn</a>{% endif %}
            </div>
        </div>
        
        <div class="date">{{ current_date }}</div>
        
        <div class="recipient">
            <div>{{ job.hiring_manager or 'Hiring Manager' }}</div>
            <div>{{ job.company }}</div>
            <div>{{ job.location }}</div>
        </div>
        
        <div class="greeting">Dear {{ job.hiring_manager or 'Hiring Manager' }},</div>
        
        <div class="content">
            <p>I am excited to apply for the {{ job.title }} position at {{ job.company }}. As a {{ user.current_title }} with {{ user.years_experience }} years of experience in {{ user.industry }}, I am eager to bring my expertise in {{ user.key_skills }} to your team.</p>
            
            <p>What draws me to {{ job.company }} is {{ job.company_appeal }}. I admire your company's {{ job.company_values }} and believe my background aligns perfectly with your mission to {{ job.company_mission }}.</p>
            
            <p>In my current role at {{ user.current_company }}, I have {{ user.current_achievements }}. This experience has honed my skills in {{ user.developed_skills }}, which I believe would transfer well to the {{ job.title }} position at {{ job.company }}.</p>
            
            <p>Some of my key accomplishments include:</p>
            <ul>
                {% for achievement in user.key_achievements %}
                <li>{{ achievement }}</li>
                {% endfor %}
            </ul>
            
            <p>I am particularly interested in the opportunity to {{ job.opportunity }} at {{ job.company }}. My background in {{ user.relevant_background }} has prepared me well for the challenges of this role, and I am confident that I can make significant contributions to your team.</p>
            
            <p>Thank you for considering my application. I would welcome the opportunity to discuss how my skills and experiences align with your needs for the {{ job.title }} position.</p>
        </div>
        
        <div class="closing">Sincerely,</div>
        
        <div class="signature">{{ user.first_name }} {{ user.last_name }}</div>
    </div>
</body>
</html>"""
                },
                'technical': {
                    'name': 'Technical Cover Letter',
                    'content': """<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{{ user.first_name }} {{ user.last_name }} - Technical Cover Letter</title>
    <style>
        body {
            font-family: 'Roboto', 'Arial', sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
        }
        .header {
            display: flex;
            justify-content: space-between;
            margin-bottom: 30px;
            padding-bottom: 20px;
            border-bottom: 1px solid #ddd;
        }
        .personal-info {
            flex: 1;
        }
        .name {
            font-size: 22px;
            font-weight: bold;
            margin-bottom: 5px;
        }
        .title {
            color: #666;
            margin-bottom: 10px;
        }
        .contact {
            font-size: 14px;
        }
        .links {
            flex: 1;
            text-align: right;
            font-size: 14px;
        }
        .date {
            margin-bottom: 30px;
        }
        .recipient {
            margin-bottom: 30px;
        }
        .greeting {
            font-size: 18px;
            margin-bottom: 20px;
        }
        .content {
            margin-bottom: 30px;
        }
        .content p {
            margin-bottom: 15px;
        }
        .tech-skills {
            background-color: #f9f9f9;
            padding: 15px;
            margin: 20px 0;
            border-left: 3px solid #ddd;
        }
        .tech-skills h3 {
            margin-top: 0;
            font-size: 16px;
        }
        .tech-skills ul {
            display: flex;
            flex-wrap: wrap;
            list-style-type: none;
            padding-left: 0;
            margin: 0;
        }
        .tech-skills li {
            margin-right: 15px;
            margin-bottom: 5px;
            background-color: #eee;
            padding: 3px 8px;
            border-radius: 3px;
            font-size: 14px;
        }
        .closing {
            margin-bottom: 20px;
        }
        .signature {
            font-weight: bold;
        }
        a {
            color: #0066cc;
            text-decoration: none;
        }
        a:hover {
            text-decoration: underline;
        }
    </style>
</head>
<body>
    <div class="header">
        <div class="personal-info">
            <div class="name">{{ user.first_name }} {{ user.last_name }}</div>
            <div class="title">{{ user.current_title }}</div>
            <div class="contact">
                {{ user.email }}<br>
                {{ user.phone }}<br>
                {{ user.location }}
            </div>
        </div>
        <div class="links">
            {% if user.linkedin_url %}<a href="{{ user.linkedin_url }}">LinkedIn</a><br>{% endif %}
            {% if user.github_url %}<a href="{{ user.github_url }}">GitHub</a><br>{% endif %}
            {% if user.portfolio_url %}<a href="{{ user.portfolio_url }}">Portfolio</a>{% endif %}
        </div>
    </div>
    
    <div class="date">{{ current_date }}</div>
    
    <div class="recipient">
        <div>{{ job.hiring_manager or 'Hiring Manager' }}</div>
        <div>{{ job.company }}</div>
        <div>{{ job.location }}</div>
    </div>
    
    <div class="greeting">Dear {{ job.hiring_manager or 'Hiring Manager' }},</div>
    
    <div class="content">
        <p>I am writing to express my interest in the {{ job.title }} position at {{ job.company }}. As a {{ user.current_title }} with expertise in {{ user.technical_expertise }}, I am excited about the opportunity to contribute to your team's technical initiatives.</p>
        
        <p>My technical background includes {{ user.years_experience }} years of experience developing {{ user.development_experience }}. In my current role at {{ user.current_company }}, I have {{ user.current_technical_achievements }}. This has strengthened my abilities in {{ user.strengthened_abilities }}, which align well with the requirements for the {{ job.title }} position.</p>
        
        <div class="tech-skills">
            <h3>Technical Skills Relevant to This Position:</h3>
            <ul>
                {% for skill in job.required_skills %}
                {% if skill in user.skills %}
                <li>{{ skill }}</li>
                {% endif %}
                {% endfor %}
            </ul>
        </div>
        
        <p>I am particularly impressed by {{ job.company }}'s {{ job.company_tech_appeal }}. Your work on {{ job.company_projects }} resonates with my interest in {{ user.technical_interests }}. I am eager to apply my experience with {{ user.relevant_technologies }} to help {{ job.company }} {{ job.company_goals }}.</p>
        
        <p>Some technical achievements that demonstrate my qualifications include:</p>
        <ul>
            {% for achievement in user.technical_achievements %}
            <li>{{ achievement }}</li>
            {% endfor %}
        </ul>
        
        <p>I am confident that my technical skills, problem-solving abilities, and passion for {{ user.passion }} would make me a valuable addition to your team. I am excited about the prospect of contributing to {{ job.company }}'s innovative work and growing professionally in this role.</p>
        
        <p>Thank you for considering my application. I look forward to discussing how my technical background and experience can benefit {{ job.company }}.</p>
    </div>
    
    <div class="closing">Sincerely,</div>
    
    <div class="signature">{{ user.first_name }} {{ user.last_name }}</div>
</body>
</html>"""
                }
            }
        
        # Write templates to files
        for template_id, template_info in templates.items():
            template_path = os.path.join(template_dir, f"{template_id}.html")
            
            try:
                with open(template_path, 'w') as f:
                    f.write(template_info['content'])
                
                logger.info(f"Created default {template_type} template: {template_id}")
            except Exception as e:
                logger.error(f"Error creating default {template_type} template {template_id}: {str(e)}")
    
    def generate_resume(self, user_profile, job_posting=None, template_id=None, output_format='html', output_path=None):
        """
        Generate a resume based on user profile and job posting.
        
        Args:
            user_profile (dict): User profile data
            job_posting (dict): Job posting data (optional)
            template_id (str): Template ID to use (optional)
            output_format (str): Output format ('html', 'pdf', 'docx')
            output_path (str): Path to save the output file (optional)
        
        Returns:
            dict: Result with generated resume information
        """
        try:
            # Select template
            if template_id and template_id in self.resume_templates:
                template = self.resume_templates[template_id]
            else:
                # Auto-select best template based on job posting
                template = self._select_best_resume_template(user_profile, job_posting)
            
            # Optimize user profile for job if provided
            if job_posting:
                optimized_profile = self._optimize_profile_for_job(user_profile, job_posting)
            else:
                optimized_profile = user_profile
            
            # Prepare template data
            template_data = {
                'user': optimized_profile,
                'job': job_posting or {},
                'current_date': datetime.now().strftime('%B %d, %Y')
            }
            
            # Render template
            jinja_template = Template(template['content'])
            rendered_content = jinja_template.render(**template_data)
            
            # Create result
            result = {
                'template_id': template['id'],
                'format': template['format'],
                'content': rendered_content
            }
            
            # Save to file if output path provided
            if output_path:
                with open(output_path, 'w') as f:
                    f.write(rendered_content)
                
                result['file_path'] = output_path
                logger.info(f"Resume saved to {output_path}")
            
            # Convert to requested format if different from template format
            if output_format != template['format']:
                converted_path = self._convert_document(rendered_content, template['format'], output_format, output_path)
                
                if converted_path:
                    result['file_path'] = converted_path
                    result['format'] = output_format
            
            return result
        except Exception as e:
            logger.error(f"Error generating resume: {str(e)}")
            return {
                'error': f"Error generating resume: {str(e)}"
            }
    
    def generate_cover_letter(self, user_profile, job_posting, template_id=None, output_format='html', output_path=None):
        """
        Generate a cover letter based on user profile and job posting.
        
        Args:
            user_profile (dict): User profile data
            job_posting (dict): Job posting data
            template_id (str): Template ID to use (optional)
            output_format (str): Output format ('html', 'pdf', 'docx')
            output_path (str): Path to save the output file (optional)
        
        Returns:
            dict: Result with generated cover letter information
        """
        try:
            # Select template
            if template_id and template_id in self.cover_letter_templates:
                template = self.cover_letter_templates[template_id]
            else:
                # Auto-select best template based on job posting
                template = self._select_best_cover_letter_template(user_profile, job_posting)
            
            # Generate tailored content for cover letter
            tailored_content = self._generate_tailored_cover_letter_content(user_profile, job_posting)
            
            # Merge tailored content with user profile
            enhanced_profile = {**user_profile, **tailored_content}
            
            # Prepare template data
            template_data = {
                'user': enhanced_profile,
                'job': job_posting,
                'current_date': datetime.now().strftime('%B %d, %Y')
            }
            
            # Render template
            jinja_template = Template(template['content'])
            rendered_content = jinja_template.render(**template_data)
            
            # Create result
            result = {
                'template_id': template['id'],
                'format': template['format'],
                'content': rendered_content
            }
            
            # Save to file if output path provided
            if output_path:
                with open(output_path, 'w') as f:
                    f.write(rendered_content)
                
                result['file_path'] = output_path
                logger.info(f"Cover letter saved to {output_path}")
            
            # Convert to requested format if different from template format
            if output_format != template['format']:
                converted_path = self._convert_document(rendered_content, template['format'], output_format, output_path)
                
                if converted_path:
                    result['file_path'] = converted_path
                    result['format'] = output_format
            
            return result
        except Exception as e:
            logger.error(f"Error generating cover letter: {str(e)}")
            return {
                'error': f"Error generating cover letter: {str(e)}"
            }
    
    def _select_best_resume_template(self, user_profile, job_posting=None):
        """
        Select the best resume template based on user profile and job posting.
        
        Args:
            user_profile (dict): User profile data
            job_posting (dict): Job posting data (optional)
        
        Returns:
            dict: Selected template
        """
        if not job_posting:
            # Default to professional template if no job posting
            return self.resume_templates.get('professional', next(iter(self.resume_templates.values())))
        
        # Check if job is technical
        technical_keywords = ['developer', 'engineer', 'programmer', 'software', 'data', 'scientist', 'analyst']
        is_technical = any(keyword in job_posting.get('title', '').lower() for keyword in technical_keywords)
        
        if is_technical:
            # Use technical template for technical jobs
            return self.resume_templates.get('technical', next(iter(self.resume_templates.values())))
        
        # Check if job is creative
        creative_keywords = ['designer', 'creative', 'artist', 'writer', 'content', 'marketing']
        is_creative = any(keyword in job_posting.get('title', '').lower() for keyword in creative_keywords)
        
        if is_creative:
            # Use modern template for creative jobs
            return self.resume_templates.get('modern', next(iter(self.resume_templates.values())))
        
        # Default to professional template
        return self.resume_templates.get('professional', next(iter(self.resume_templates.values())))
    
    def _select_best_cover_letter_template(self, user_profile, job_posting):
        """
        Select the best cover letter template based on user profile and job posting.
        
        Args:
            user_profile (dict): User profile data
            job_posting (dict): Job posting data
        
        Returns:
            dict: Selected template
        """
        # Check if job is technical
        technical_keywords = ['developer', 'engineer', 'programmer', 'software', 'data', 'scientist', 'analyst']
        is_technical = any(keyword in job_posting.get('title', '').lower() for keyword in technical_keywords)
        
        if is_technical:
            # Use technical template for technical jobs
            return self.cover_letter_templates.get('technical', next(iter(self.cover_letter_templates.values())))
        
        # Check if job is creative
        creative_keywords = ['designer', 'creative', 'artist', 'writer', 'content', 'marketing']
        is_creative = any(keyword in job_posting.get('title', '').lower() for keyword in creative_keywords)
        
        if is_creative:
            # Use modern template for creative jobs
            return self.cover_letter_templates.get('modern', next(iter(self.cover_letter_templates.values())))
        
        # Default to professional template
        return self.cover_letter_templates.get('professional', next(iter(self.cover_letter_templates.values())))
    
    def _optimize_profile_for_job(self, user_profile, job_posting):
        """
        Optimize user profile for a specific job posting.
        
        Args:
            user_profile (dict): User profile data
            job_posting (dict): Job posting data
        
        Returns:
            dict: Optimized user profile
        """
        # Create a copy of the user profile to avoid modifying the original
        optimized_profile = user_profile.copy()
        
        # Extract job requirements
        job_skills = self._extract_skills_from_job(job_posting)
        job_title = job_posting.get('title', '').lower()
        job_description = job_posting.get('description', '').lower()
        
        # Prioritize matching skills
        if 'skills' in optimized_profile and job_skills:
            user_skills = optimized_profile['skills']
            
            # Move matching skills to the front
            matching_skills = [skill for skill in user_skills if skill.lower() in [s.lower() for s in job_skills]]
            other_skills = [skill for skill in user_skills if skill.lower() not in [s.lower() for s in job_skills]]
            
            optimized_profile['skills'] = matching_skills + other_skills
        
        # Prioritize relevant work experience
        if 'work_history' in optimized_profile and isinstance(optimized_profile['work_history'], list):
            for job in optimized_profile['work_history']:
                if 'achievements' in job and isinstance(job['achievements'], list):
                    # Prioritize achievements that match job requirements
                    relevant_achievements = []
                    other_achievements = []
                    
                    for achievement in job['achievements']:
                        # Check if achievement is relevant to job
                        is_relevant = False
                        
                        # Check for job title keywords
                        title_keywords = job_title.split()
                        if any(keyword in achievement.lower() for keyword in title_keywords if len(keyword) > 3):
                            is_relevant = True
                        
                        # Check for job skills
                        if any(skill.lower() in achievement.lower() for skill in job_skills):
                            is_relevant = True
                        
                        if is_relevant:
                            relevant_achievements.append(achievement)
                        else:
                            other_achievements.append(achievement)
                    
                    # Update achievements with prioritized list
                    job['achievements'] = relevant_achievements + other_achievements
        
        # Customize summary if present
        if 'summary' in optimized_profile and job_posting.get('title') and job_posting.get('company'):
            # Extract years of experience
            years_experience = self._extract_years_experience(optimized_profile)
            
            # Create tailored summary
            original_summary = optimized_profile['summary']
            tailored_summary = f"Experienced {job_posting['title']} with {years_experience}+ years of expertise in {', '.join(job_skills[:3] if job_skills else ['relevant skills'])}. {original_summary}"
            
            optimized_profile['summary'] = tailored_summary
        
        return optimized_profile
    
    def _extract_skills_from_job(self, job_posting):
        """
        Extract skills from job posting.
        
        Args:
            job_posting (dict): Job posting data
        
        Returns:
            list: Extracted skills
        """
        skills = []
        
        # Extract from skills field
        if 'required_skills' in job_posting:
            if isinstance(job_posting['required_skills'], list):
                skills.extend(job_posting['required_skills'])
            elif isinstance(job_posting['required_skills'], str):
                skills.extend([s.strip() for s in job_posting['required_skills'].split(',')])
        
        # Extract from description
        if 'description' in job_posting:
            # This is a simplified approach; in a real system, you'd use NLP techniques
            description = job_posting['description'].lower()
            common_skills = [
                'python', 'javascript', 'java', 'c++', 'c#', 'react', 'angular', 'vue',
                'node.js', 'django', 'flask', 'spring', 'sql', 'nosql', 'mongodb',
                'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'ci/cd', 'git',
                'agile', 'scrum', 'product management', 'project management',
                'machine learning', 'ai', 'data science', 'data analysis',
                'full stack', 'frontend', 'backend', 'devops', 'cloud', 'security'
            ]
            
            for skill in common_skills:
                if skill in description:
                    skills.append(skill)
        
        # Remove duplicates and return
        return list(set(skills))
    
    def _extract_years_experience(self, user_profile):
        """
        Extract years of experience from user profile.
        
        Args:
            user_profile (dict): User profile data
        
        Returns:
            int: Years of experience
        """
        years = 0
        
        # Extract from work_history
        if 'work_history' in user_profile and isinstance(user_profile['work_history'], list):
            for job in user_profile['work_history']:
                if isinstance(job, dict) and 'duration' in job:
                    if isinstance(job['duration'], (int, float)):
                        years += job['duration']
                    elif isinstance(job['duration'], str):
                        # Try to extract years from string
                        match = re.search(r'(\d+)', job['duration'])
                        if match:
                            years += int(match.group(1))
        
        # Extract from experience field
        if 'experience' in user_profile:
            if isinstance(user_profile['experience'], (int, float)):
                years = max(years, user_profile['experience'])
            elif isinstance(user_profile['experience'], str):
                # Try to extract years from string
                match = re.search(r'(\d+)', user_profile['experience'])
                if match:
                    years = max(years, int(match.group(1)))
        
        return max(1, int(years))  # Minimum 1 year
    
    def _generate_tailored_cover_letter_content(self, user_profile, job_posting):
        """
        Generate tailored content for cover letter.
        
        Args:
            user_profile (dict): User profile data
            job_posting (dict): Job posting data
        
        Returns:
            dict: Tailored content
        """
        # Extract key information
        job_title = job_posting.get('title', '')
        company = job_posting.get('company', '')
        job_skills = self._extract_skills_from_job(job_posting)
        user_skills = user_profile.get('skills', [])
        
        # Find matching skills
        matching_skills = [skill for skill in user_skills if skill.lower() in [s.lower() for s in job_skills]]
        
        # Extract years of experience
        years_experience = self._extract_years_experience(user_profile)
        
        # Determine user's industry
        industry = self._determine_industry(user_profile, job_posting)
        
        # Extract current company and title
        current_company = ''
        current_title = ''
        if 'work_history' in user_profile and isinstance(user_profile['work_history'], list) and user_profile['work_history']:
            current_job = user_profile['work_history'][0]
            current_company = current_job.get('company', '')
            current_title = current_job.get('title', '')
        
        # Generate tailored content
        tailored_content = {
            'years_experience': years_experience,
            'industry': industry,
            'key_skills': ', '.join(matching_skills[:3]) if matching_skills else 'relevant skills',
            'current_company': current_company,
            'current_title': current_title,
            'background': industry,
            'experience': ', '.join(matching_skills[:3]) if matching_skills else 'relevant areas',
            'company_appeal': f"your reputation for {self._generate_company_appeal(company, job_posting)}",
            'position_appeal': f"it allows me to leverage my expertise in {', '.join(matching_skills[:3]) if matching_skills else 'relevant skills'}",
            'opportunity': f"contribute to {company}'s success through {self._generate_opportunity(job_title, job_posting)}"
        }
        
        # Generate achievements if not present
        if 'key_achievements' not in user_profile or not user_profile['key_achievements']:
            tailored_content['key_achievements'] = self._generate_achievements(user_profile, job_posting)
        
        # Generate technical content if job is technical
        technical_keywords = ['developer', 'engineer', 'programmer', 'software', 'data', 'scientist', 'analyst']
        is_technical = any(keyword in job_title.lower() for keyword in technical_keywords)
        
        if is_technical:
            tailored_content.update({
                'technical_expertise': ', '.join(matching_skills[:3]) if matching_skills else 'software development',
                'development_experience': self._generate_development_experience(user_profile, job_posting),
                'technical_interests': self._generate_technical_interests(user_profile, job_posting),
                'relevant_technologies': ', '.join(matching_skills[:3]) if matching_skills else 'relevant technologies',
                'passion': 'technology and innovation',
                'technical_achievements': self._generate_technical_achievements(user_profile, job_posting)
            })
        
        return tailored_content
    
    def _determine_industry(self, user_profile, job_posting):
        """
        Determine user's industry based on profile and job posting.
        
        Args:
            user_profile (dict): User profile data
            job_posting (dict): Job posting data
        
        Returns:
            str: Industry
        """
        # Check if industry is explicitly specified
        if 'industry' in user_profile:
            return user_profile['industry']
        
        # Try to determine from work history
        if 'work_history' in user_profile and isinstance(user_profile['work_history'], list) and user_profile['work_history']:
            # Check job titles for industry clues
            tech_keywords = ['developer', 'engineer', 'programmer', 'software', 'data', 'IT']
            finance_keywords = ['finance', 'accounting', 'financial', 'banker', 'investment']
            marketing_keywords = ['marketing', 'advertiser', 'brand', 'content', 'social media']
            
            for job in user_profile['work_history']:
                title = job.get('title', '').lower()
                
                if any(keyword.lower() in title for keyword in tech_keywords):
                    return 'technology'
                elif any(keyword.lower() in title for keyword in finance_keywords):
                    return 'finance'
                elif any(keyword.lower() in title for keyword in marketing_keywords):
                    return 'marketing'
        
        # Default to industry based on job posting
        job_title = job_posting.get('title', '').lower()
        
        if any(keyword in job_title for keyword in ['developer', 'engineer', 'programmer', 'software']):
            return 'software development'
        elif any(keyword in job_title for keyword in ['data', 'analyst', 'scientist']):
            return 'data science'
        elif any(keyword in job_title for keyword in ['designer', 'ux', 'ui']):
            return 'design'
        elif any(keyword in job_title for keyword in ['manager', 'director', 'lead']):
            return 'management'
        
        return 'professional services'
    
    def _generate_company_appeal(self, company, job_posting):
        """
        Generate company appeal statement.
        
        Args:
            company (str): Company name
            job_posting (dict): Job posting data
        
        Returns:
            str: Company appeal statement
        """
        # List of possible appeals
        appeals = [
            "innovation and forward-thinking approach",
            "commitment to excellence and quality",
            "collaborative and inclusive culture",
            "industry leadership and vision",
            "focus on customer satisfaction and service",
            "dedication to professional development and growth",
            "cutting-edge products and services",
            "strong values and ethical practices"
        ]
        
        # Try to extract company values from job posting
        company_values = []
        if 'description' in job_posting:
            description = job_posting['description'].lower()
            
            value_keywords = {
                'innovation': ['innovation', 'innovative', 'cutting-edge', 'leading'],
                'collaboration': ['collaboration', 'collaborative', 'team', 'teamwork'],
                'excellence': ['excellence', 'quality', 'best', 'exceptional'],
                'diversity': ['diversity', 'inclusion', 'inclusive', 'diverse'],
                'growth': ['growth', 'development', 'learning', 'opportunity']
            }
            
            for value, keywords in value_keywords.items():
                if any(keyword in description for keyword in keywords):
                    company_values.append(value)
        
        if company_values:
            return f"{', '.join(company_values[:-1]) + ' and ' + company_values[-1] if len(company_values) > 1 else company_values[0]}"
        
        # If no values found, return random appeal
        import random
        return random.choice(appeals)
    
    def _generate_opportunity(self, job_title, job_posting):
        """
        Generate opportunity statement.
        
        Args:
            job_title (str): Job title
            job_posting (dict): Job posting data
        
        Returns:
            str: Opportunity statement
        """
        # Extract responsibilities from job posting
        responsibilities = []
        if 'description' in job_posting:
            description = job_posting['description'].lower()
            
            # Look for responsibility section
            resp_section = re.search(r'responsibilities?:(.*?)(?:requirements|qualifications|about you|about the role|about the company|about us)', description, re.IGNORECASE | re.DOTALL)
            
            if resp_section:
                resp_text = resp_section.group(1)
                
                # Extract bullet points
                bullet_points = re.findall(r'[•\-*]\s*(.*?)(?=[•\-*]|\Z)', resp_text, re.DOTALL)
                
                if bullet_points:
                    responsibilities = [point.strip() for point in bullet_points if point.strip()]
        
        if responsibilities:
            # Pick a few key responsibilities
            import random
            selected_resps = random.sample(responsibilities, min(2, len(responsibilities)))
            
            # Convert to gerund form (ing)
            gerund_resps = []
            for resp in selected_resps:
                # Extract first verb
                verb_match = re.match(r'(\w+)', resp)
                if verb_match:
                    verb = verb_match.group(1).lower()
                    
                    # Convert to gerund
                    if verb.endswith('e'):
                        gerund = verb[:-1] + 'ing'
                    elif verb.endswith('y'):
                        gerund = verb[:-1] + 'ying'
                    else:
                        gerund = verb + 'ing'
                    
                    # Replace verb with gerund
                    gerund_resp = re.sub(r'^\w+', gerund, resp, 1)
                    gerund_resps.append(gerund_resp)
                else:
                    gerund_resps.append(resp)
            
            return ' and '.join(gerund_resps)
        
        # Default opportunities based on job title
        if 'developer' in job_title.lower() or 'engineer' in job_title.lower():
            return "developing innovative solutions and contributing to technical excellence"
        elif 'manager' in job_title.lower() or 'director' in job_title.lower():
            return "leading teams and driving strategic initiatives"
        elif 'analyst' in job_title.lower() or 'scientist' in job_title.lower():
            return "analyzing data and delivering actionable insights"
        elif 'designer' in job_title.lower():
            return "creating engaging user experiences and innovative designs"
        
        return "contributing my skills and experience to your team's success"
    
    def _generate_achievements(self, user_profile, job_posting):
        """
        Generate achievements based on user profile and job posting.
        
        Args:
            user_profile (dict): User profile data
            job_posting (dict): Job posting data
        
        Returns:
            list: Generated achievements
        """
        achievements = []
        
        # Extract work history
        work_history = user_profile.get('work_history', [])
        
        # Extract skills
        user_skills = user_profile.get('skills', [])
        job_skills = self._extract_skills_from_job(job_posting)
        
        # Generate achievements based on work history
        if work_history and isinstance(work_history, list):
            for job in work_history[:2]:  # Use only the two most recent jobs
                if isinstance(job, dict):
                    title = job.get('title', '')
                    company = job.get('company', '')
                    
                    if title and company:
                        # Generate achievement based on job title
                        if 'developer' in title.lower() or 'engineer' in title.lower():
                            achievements.append(f"Developed and implemented {self._random_technical_feature()} at {company}, resulting in {self._random_positive_outcome()}")
                        elif 'manager' in title.lower() or 'lead' in title.lower():
                            achievements.append(f"Led a team of {self._random_team_size()} at {company}, delivering {self._random_project_outcome()}")
                        elif 'analyst' in title.lower() or 'scientist' in title.lower():
                            achievements.append(f"Analyzed {self._random_data_source()} at {company}, providing insights that {self._random_analysis_outcome()}")
                        else:
                            achievements.append(f"Successfully {self._random_professional_achievement()} at {company}, which {self._random_positive_outcome()}")
        
        # Generate achievements based on matching skills
        matching_skills = [skill for skill in user_skills if skill.lower() in [s.lower() for s in job_skills]]
        
        if matching_skills:
            for skill in matching_skills[:2]:  # Use only the two most relevant skills
                achievements.append(f"Utilized expertise in {skill} to {self._random_skill_achievement()}")
        
        # Add generic achievements if needed
        while len(achievements) < 3:
            achievements.append(self._random_generic_achievement())
        
        return achievements
    
    def _generate_development_experience(self, user_profile, job_posting):
        """
        Generate development experience statement.
        
        Args:
            user_profile (dict): User profile data
            job_posting (dict): Job posting data
        
        Returns:
            str: Development experience statement
        """
        # Extract skills
        user_skills = user_profile.get('skills', [])
        job_skills = self._extract_skills_from_job(job_posting)
        
        # Find matching skills
        matching_skills = [skill for skill in user_skills if skill.lower() in [s.lower() for s in job_skills]]
        
        if matching_skills:
            # Group skills by category
            frontend_skills = [skill for skill in matching_skills if skill.lower() in ['react', 'angular', 'vue', 'javascript', 'html', 'css']]
            backend_skills = [skill for skill in matching_skills if skill.lower() in ['python', 'java', 'node.js', 'django', 'flask', 'spring', 'express']]
            database_skills = [skill for skill in matching_skills if skill.lower() in ['sql', 'mysql', 'postgresql', 'mongodb', 'nosql']]
            cloud_skills = [skill for skill in matching_skills if skill.lower() in ['aws', 'azure', 'gcp', 'cloud', 'docker', 'kubernetes']]
            
            experience_parts = []
            
            if frontend_skills:
                experience_parts.append(f"frontend applications using {', '.join(frontend_skills[:3])}")
            
            if backend_skills:
                experience_parts.append(f"backend systems with {', '.join(backend_skills[:3])}")
            
            if database_skills:
                experience_parts.append(f"database solutions including {', '.join(database_skills[:3])}")
            
            if cloud_skills:
                experience_parts.append(f"cloud infrastructure on {', '.join(cloud_skills[:3])}")
            
            if experience_parts:
                if len(experience_parts) > 1:
                    return f"{' and '.join([', '.join(experience_parts[:-1]), experience_parts[-1]])}"
                else:
                    return experience_parts[0]
        
        # Default statement
        return "software applications and systems using various programming languages and frameworks"
    
    def _generate_technical_interests(self, user_profile, job_posting):
        """
        Generate technical interests statement.
        
        Args:
            user_profile (dict): User profile data
            job_posting (dict): Job posting data
        
        Returns:
            str: Technical interests statement
        """
        # Extract job title and description
        job_title = job_posting.get('title', '').lower()
        job_description = job_posting.get('description', '').lower()
        
        # Check for specific technical areas
        if 'data' in job_title or 'data' in job_description:
            return "data analysis and machine learning"
        elif 'cloud' in job_title or 'cloud' in job_description:
            return "cloud architecture and scalable systems"
        elif 'frontend' in job_title or 'frontend' in job_description or 'front-end' in job_title or 'front-end' in job_description:
            return "creating intuitive and responsive user interfaces"
        elif 'backend' in job_title or 'backend' in job_description or 'back-end' in job_title or 'back-end' in job_description:
            return "building robust and scalable backend systems"
        elif 'fullstack' in job_title or 'fullstack' in job_description or 'full-stack' in job_title or 'full-stack' in job_description:
            return "developing end-to-end solutions"
        elif 'mobile' in job_title or 'mobile' in job_description:
            return "mobile application development"
        elif 'security' in job_title or 'security' in job_description:
            return "cybersecurity and secure coding practices"
        elif 'devops' in job_title or 'devops' in job_description:
            return "DevOps practices and CI/CD pipelines"
        
        # Default interest
        return "solving complex technical challenges and continuous learning"
    
    def _generate_technical_achievements(self, user_profile, job_posting):
        """
        Generate technical achievements.
        
        Args:
            user_profile (dict): User profile data
            job_posting (dict): Job posting data
        
        Returns:
            list: Technical achievements
        """
        achievements = []
        
        # Extract skills
        user_skills = user_profile.get('skills', [])
        job_skills = self._extract_skills_from_job(job_posting)
        
        # Find matching skills
        matching_skills = [skill for skill in user_skills if skill.lower() in [s.lower() for s in job_skills]]
        
        # Generate achievements based on matching skills
        if matching_skills:
            for skill in matching_skills[:3]:  # Use only the three most relevant skills
                achievements.append(f"Implemented {self._random_technical_feature()} using {skill}, resulting in {self._random_technical_outcome()}")
        
        # Add generic technical achievements if needed
        while len(achievements) < 3:
            achievements.append(self._random_technical_achievement())
        
        return achievements
    
    def _random_technical_feature(self):
        """Generate a random technical feature."""
        features = [
            "a RESTful API",
            "a microservices architecture",
            "a responsive web interface",
            "a data processing pipeline",
            "a real-time analytics dashboard",
            "a cloud-based deployment system",
            "a database optimization strategy",
            "an automated testing framework",
            "a continuous integration pipeline",
            "a secure authentication system"
        ]
        import random
        return random.choice(features)
    
    def _random_positive_outcome(self):
        """Generate a random positive outcome."""
        outcomes = [
            "increased efficiency by 30%",
            "reduced costs by 25%",
            "improved customer satisfaction by 40%",
            "accelerated time-to-market by 35%",
            "enhanced system reliability by 50%",
            "boosted team productivity by 20%",
            "generated $500K in additional revenue",
            "decreased error rates by 45%",
            "streamlined operations significantly",
            "received excellent stakeholder feedback"
        ]
        import random
        return random.choice(outcomes)
    
    def _random_team_size(self):
        """Generate a random team size."""
        import random
        return f"{random.randint(3, 15)} professionals"
    
    def _random_project_outcome(self):
        """Generate a random project outcome."""
        outcomes = [
            "projects on time and under budget",
            "award-winning solutions",
            "innovative products that exceeded expectations",
            "strategic initiatives with measurable business impact",
            "successful digital transformation efforts",
            "high-performance systems with 99.9% uptime",
            "enterprise solutions adopted across the organization",
            "customer-focused applications with excellent reviews"
        ]
        import random
        return random.choice(outcomes)
    
    def _random_data_source(self):
        """Generate a random data source."""
        sources = [
            "customer behavior data",
            "market trends",
            "operational metrics",
            "financial performance indicators",
            "user engagement statistics",
            "competitive intelligence",
            "supply chain data",
            "sales and marketing analytics",
            "product performance metrics",
            "large datasets from multiple sources"
        ]
        import random
        return random.choice(sources)
    
    def _random_analysis_outcome(self):
        """Generate a random analysis outcome."""
        outcomes = [
            "drove strategic decision-making",
            "identified $1M in cost-saving opportunities",
            "revealed new market opportunities",
            "optimized resource allocation",
            "improved forecasting accuracy by 40%",
            "enabled data-driven product improvements",
            "supported successful business expansion",
            "highlighted critical performance issues",
            "informed successful marketing campaigns",
            "contributed to 25% revenue growth"
        ]
        import random
        return random.choice(outcomes)
    
    def _random_professional_achievement(self):
        """Generate a random professional achievement."""
        achievements = [
            "managed key client relationships",
            "redesigned critical business processes",
            "launched innovative initiatives",
            "implemented best practices",
            "coordinated cross-functional projects",
            "developed strategic partnerships",
            "created comprehensive documentation",
            "established new quality standards",
            "facilitated successful knowledge transfer",
            "optimized workflow procedures"
        ]
        import random
        return random.choice(achievements)
    
    def _random_skill_achievement(self):
        """Generate a random skill achievement."""
        achievements = [
            "solve complex business challenges efficiently",
            "develop innovative solutions that increased ROI by 30%",
            "streamline critical processes and reduce overhead",
            "create scalable systems that supported business growth",
            "implement best practices that became company standards",
            "deliver high-quality results under tight deadlines",
            "mentor team members and improve overall capabilities",
            "optimize performance in mission-critical applications",
            "reduce technical debt while adding new features",
            "integrate disparate systems into a cohesive solution"
        ]
        import random
        return random.choice(achievements)
    
    def _random_generic_achievement(self):
        """Generate a random generic achievement."""
        achievements = [
            "Consistently exceeded performance targets and received recognition for excellence",
            "Successfully managed multiple priorities in fast-paced environments while maintaining quality",
            "Collaborated effectively with cross-functional teams to deliver integrated solutions",
            "Identified and implemented process improvements that enhanced efficiency by 20%",
            "Recognized for problem-solving abilities and innovative approaches to challenges",
            "Adapted quickly to changing requirements and delivered results under pressure",
            "Maintained strong relationships with stakeholders at all levels of the organization",
            "Contributed to strategic planning that aligned with organizational objectives",
            "Demonstrated leadership in challenging situations and guided teams to successful outcomes",
            "Received consistently positive feedback from managers, peers, and clients"
        ]
        import random
        return random.choice(achievements)
    
    def _random_technical_outcome(self):
        """Generate a random technical outcome."""
        outcomes = [
            "a 40% improvement in system performance",
            "95% test coverage and significantly reduced bugs",
            "successful deployment to thousands of users",
            "a 50% reduction in page load times",
            "seamless integration with existing systems",
            "enhanced security and compliance with industry standards",
            "a scalable solution that handled 3x the previous load",
            "decreased server costs by 35%",
            "positive user feedback and increased engagement",
            "a maintainable codebase that facilitated future enhancements"
        ]
        import random
        return random.choice(outcomes)
    
    def _random_technical_achievement(self):
        """Generate a random technical achievement."""
        achievements = [
            "Architected and implemented a scalable cloud infrastructure that reduced hosting costs by 30%",
            "Developed a CI/CD pipeline that decreased deployment time from days to minutes",
            "Optimized database queries resulting in 50% faster application response times",
            "Created a modular component library that accelerated development across multiple projects",
            "Implemented automated testing strategies that caught 95% of bugs before production",
            "Refactored legacy codebase, reducing complexity and improving maintainability",
            "Designed and built RESTful APIs that enabled seamless integration with partner systems",
            "Led the migration from monolithic architecture to microservices, improving system resilience",
            "Developed data processing algorithms that handled terabytes of information efficiently",
            "Implemented security best practices that protected sensitive data and prevented breaches"
        ]
        import random
        return random.choice(achievements)
    
    def _convert_document(self, content, source_format, target_format, output_path=None):
        """
        Convert document from one format to another.
        
        Args:
            content (str): Document content
            source_format (str): Source format
            target_format (str): Target format
            output_path (str): Output path (optional)
        
        Returns:
            str: Path to converted document or None if conversion failed
        """
        try:
            # Determine output path if not provided
            if not output_path:
                import tempfile
                output_dir = tempfile.gettempdir()
                output_filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{target_format}"
                output_path = os.path.join(output_dir, output_filename)
            
            # If output path has wrong extension, fix it
            if not output_path.endswith(f".{target_format}"):
                output_path = f"{os.path.splitext(output_path)[0]}.{target_format}"
            
            # Create temporary file for source content
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=f".{source_format}", delete=False) as temp:
                temp_path = temp.name
                temp.write(content.encode('utf-8'))
            
            # Convert based on formats
            if source_format == 'html' and target_format == 'pdf':
                self._convert_html_to_pdf(temp_path, output_path)
            elif source_format == 'html' and target_format == 'docx':
                self._convert_html_to_docx(temp_path, output_path)
            else:
                logger.warning(f"Conversion from {source_format} to {target_format} not supported")
                return None
            
            # Clean up temporary file
            os.unlink(temp_path)
            
            return output_path
        except Exception as e:
            logger.error(f"Error converting document: {str(e)}")
            return None
    
    def _convert_html_to_pdf(self, html_path, pdf_path):
        """
        Convert HTML to PDF.
        
        Args:
            html_path (str): Path to HTML file
            pdf_path (str): Path to output PDF file
        """
        try:
            from weasyprint import HTML
            HTML(html_path).write_pdf(pdf_path)
            logger.info(f"Converted HTML to PDF: {pdf_path}")
        except Exception as e:
            logger.error(f"Error converting HTML to PDF: {str(e)}")
            raise
    
    def _convert_html_to_docx(self, html_path, docx_path):
        """
        Convert HTML to DOCX.
        
        Args:
            html_path (str): Path to HTML file
            docx_path (str): Path to output DOCX file
        """
        try:
            # This is a simplified approach; in a real system, you'd use a library like python-docx
            # For now, we'll just create a simple DOCX with the HTML content
            from docx import Document
            from bs4 import BeautifulSoup
            
            # Read HTML content
            with open(html_path, 'r') as f:
                html_content = f.read()
            
            # Parse HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Create DOCX
            doc = Document()
            
            # Add title
            title = soup.title.string if soup.title else "Document"
            doc.add_heading(title, 0)
            
            # Add content (simplified)
            for p in soup.find_all('p'):
                doc.add_paragraph(p.get_text())
            
            # Add headings
            for i in range(1, 7):
                for h in soup.find_all(f'h{i}'):
                    doc.add_heading(h.get_text(), i)
            
            # Save DOCX
            doc.save(docx_path)
            logger.info(f"Converted HTML to DOCX: {docx_path}")
        except Exception as e:
            logger.error(f"Error converting HTML to DOCX: {str(e)}")
            raise
